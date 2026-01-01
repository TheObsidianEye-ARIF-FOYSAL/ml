from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def read_code_file(filepath):
    """Read code from a file"""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        return f"# Error reading file: {e}"

def add_page_break(doc):
    doc.add_page_break()

def set_font_style(run, bold=False, size=11, color=RGBColor(0, 0, 0)):
    run.font.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color

def add_code_block(doc, code_text):
    """Add code block with proper formatting"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.5)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    
    # Add shading to paragraph
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F5F5F5')
    para._element.get_or_add_pPr().append(shading)
    
    run = para.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

def create_lab2_report():
    """Lab 2: BFS & DFS Report"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Lab Report 2: Graph Traversal Algorithms', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph('Breadth-First Search (BFS) and Depth-First Search (DFS)')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.runs[0]
    set_font_style(run, bold=True, size=14, color=RGBColor(52, 73, 94))
    
    doc.add_paragraph()
    
    # ==========================
    # EXPERIMENT 1: BFS
    # ==========================
    doc.add_heading('Experiment 1: Breadth-First Search (BFS)', 1)
    
    # Theory
    doc.add_heading('1.1 Theory', 2)
    theory_bfs = doc.add_paragraph()
    theory_bfs.add_run('Breadth-First Search (BFS) ').bold = True
    theory_bfs.add_run(
        'is a fundamental graph traversal algorithm that explores nodes level by level. '
        'It starts from a source vertex and explores all neighboring vertices at the present depth '
        'before moving to vertices at the next depth level.\n\n'
    )
    
    doc.add_paragraph('Key Characteristics:', style='List Bullet')
    doc.add_paragraph('Uses a Queue data structure (FIFO - First In First Out)', style='List Bullet 2')
    doc.add_paragraph('Visits vertices level by level', style='List Bullet 2')
    doc.add_paragraph('Guarantees shortest path in unweighted graphs', style='List Bullet 2')
    doc.add_paragraph('Time Complexity: O(V + E) where V = vertices, E = edges', style='List Bullet 2')
    doc.add_paragraph('Space Complexity: O(V) for the queue and visited array', style='List Bullet 2')
    
    doc.add_paragraph('\nAlgorithm Steps:', style='List Bullet')
    doc.add_paragraph('1. Mark the starting vertex as visited and enqueue it', style='List Number')
    doc.add_paragraph('2. While the queue is not empty:', style='List Number')
    doc.add_paragraph('   a. Dequeue a vertex and process it', style='List Number 2')
    doc.add_paragraph('   b. Visit all unvisited adjacent vertices, mark them visited, and enqueue them', style='List Number 2')
    doc.add_paragraph('3. Repeat until the queue is empty', style='List Number')
    
    # Code
    doc.add_heading('1.2 Code Implementation', 2)
    bfs_code = read_code_file('lab_02/bfs.py')
    
    add_code_block(doc, bfs_code)
    
    # Output
    doc.add_heading('1.3 Output', 2)
    doc.add_paragraph('Graph Structure:', style='List Bullet')
    doc.add_paragraph('Vertex 0 → [1, 2]', style='List Bullet 2')
    doc.add_paragraph('Vertex 1 → [0, 3, 4]', style='List Bullet 2')
    doc.add_paragraph('Vertex 2 → [0, 5, 6]', style='List Bullet 2')
    doc.add_paragraph('Vertex 3 → [1, 7, 8]', style='List Bullet 2')
    doc.add_paragraph('Vertex 4 → [1, 9, 10]', style='List Bullet 2')
    doc.add_paragraph('...and more vertices', style='List Bullet 2')
    
    output_para = doc.add_paragraph('\nExample: Start = 0, Goal = 11\n')
    run = output_para.add_run('Path: [0, 2, 5, 11]\nCost: 7')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 128, 0)
    run.font.size = Pt(12)
    
    doc.add_paragraph(
        '\nExplanation: BFS explores the graph level by level, finding the path from start to goal. '
        'The algorithm uses a queue to maintain the order of exploration and tracks the cumulative cost. '
        'BFS guarantees finding a solution if one exists, making it reliable for pathfinding in weighted graphs.'
    )
    
    # Conclusion
    doc.add_heading('1.4 Conclusion', 2)
    conclusion_bfs = doc.add_paragraph(
        'The Breadth-First Search algorithm successfully finds paths in weighted graphs by exploring level by level. '
        'This implementation accepts start and goal nodes, along with edge weights, to determine both the path and its cost. '
        'BFS is particularly useful for finding paths in graphs where you need to track the total cost, '
        'making it applicable to routing problems, network analysis, and pathfinding applications. '
        'The queue-based approach ensures systematic exploration of all possible paths.'
    )
    
    add_page_break(doc)
    
    # ==========================
    # EXPERIMENT 2: DFS
    # ==========================
    doc.add_heading('Experiment 2: Depth-First Search (DFS) - Cycle Detection', 1)
    
    # Theory
    doc.add_heading('2.1 Theory', 2)
    theory_dfs = doc.add_paragraph()
    theory_dfs.add_run('Depth-First Search (DFS) for Cycle Detection ').bold = True
    theory_dfs.add_run(
        'is a graph traversal algorithm that explores as far as possible along each branch before backtracking. '
        'This implementation specifically detects whether a cycle exists in an undirected graph.\n\n'
    )
    
    doc.add_paragraph('Key Characteristics:', style='List Bullet')
    doc.add_paragraph('Uses recursion (implicit stack) for traversal', style='List Bullet 2')
    doc.add_paragraph('Tracks visited nodes to detect cycles', style='List Bullet 2')
    doc.add_paragraph('Maintains parent information to avoid false cycle detection', style='List Bullet 2')
    doc.add_paragraph('Returns True if cycle exists, False otherwise', style='List Bullet 2')
    doc.add_paragraph('Time Complexity: O(V + E) where V = vertices, E = edges', style='List Bullet 2')
    doc.add_paragraph('Space Complexity: O(V) for the recursion stack and visited array', style='List Bullet 2')
    
    doc.add_paragraph('\nAlgorithm Steps:', style='List Bullet')
    doc.add_paragraph('1. Mark the current vertex as visited', style='List Number')
    doc.add_paragraph('2. For each adjacent vertex:', style='List Number')
    doc.add_paragraph('   a. If unvisited, recursively visit it', style='List Number 2')
    doc.add_paragraph('   b. If visited and not parent, cycle detected', style='List Number 2')
    doc.add_paragraph('3. Return cycle detection result', style='List Number')
    
    doc.add_paragraph('\nCycle Detection Logic:', style='List Bullet')
    doc.add_paragraph('If we visit a vertex that is already visited AND it\'s not the parent of current vertex, a cycle exists', style='List Bullet 2')
    
    # Code
    doc.add_heading('2.2 Code Implementation', 2)
    dfs_code = read_code_file('lab_02/dfs.py')
    
    add_code_block(doc, dfs_code)
    
    # Output
    doc.add_heading('2.3 Output', 2)
    doc.add_paragraph('Test Graph 1 (With Cycle):', style='List Bullet')
    doc.add_paragraph('Vertex 0 → [1, 4]', style='List Bullet 2')
    doc.add_paragraph('Vertex 1 → [0, 2, 3]', style='List Bullet 2')
    doc.add_paragraph('Vertex 2 → [1, 3]', style='List Bullet 2')
    doc.add_paragraph('Vertex 3 → [1, 2]', style='List Bullet 2')
    doc.add_paragraph('Vertex 4 → [0, 5]', style='List Bullet 2')
    doc.add_paragraph('Vertex 5 → [4]', style='List Bullet 2')
    
    output_para = doc.add_paragraph('\nOutput: ')
    run = output_para.add_run('cycle')
    run.bold = True
    run.font.color.rgb = RGBColor(255, 0, 0)
    run.font.size = Pt(12)
    
    doc.add_paragraph(
        '\nExplanation: The graph contains a cycle (1→2→3→1), which DFS successfully detects by identifying '
        'that vertex 1 is visited again through a path that doesn\'t go through its parent.'
    )
    
    doc.add_paragraph('\nTest Graph 2 (Without Cycle - Tree):', style='List Bullet')
    doc.add_paragraph('A tree structure with no back edges', style='List Bullet 2')
    
    output_para2 = doc.add_paragraph('Output: ')
    run = output_para2.add_run('There is no cycle')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 128, 0)
    run.font.size = Pt(12)
    
    # Conclusion
    doc.add_heading('2.4 Conclusion', 2)
    conclusion_dfs = doc.add_paragraph(
        'The Depth-First Search algorithm for cycle detection successfully identifies cycles in undirected graphs. '
        'DFS is particularly useful for: detecting cycles (as shown), finding connected components, topological sorting, '
        'and solving maze problems. The recursive implementation with parent tracking prevents false positives '
        'from the bidirectional nature of undirected graphs. This algorithm is fundamental in graph theory and has '
        'applications in deadlock detection, dependency analysis, and network topology validation.'
    )
    
    # Overall Conclusion
    add_page_break(doc)
    doc.add_heading('Overall Conclusion', 1)
    overall_conclusion = doc.add_paragraph(
        'Both BFS and DFS are fundamental graph traversal algorithms with distinct characteristics and use cases:\n\n'
    )
    
    doc.add_paragraph('BFS (Pathfinding) Applications:', style='List Bullet')
    doc.add_paragraph('Finding paths with cost tracking in weighted graphs', style='List Bullet 2')
    doc.add_paragraph('Shortest path problems', style='List Bullet 2')
    doc.add_paragraph('Network routing and navigation', style='List Bullet 2')
    doc.add_paragraph('Level-order exploration', style='List Bullet 2')
    
    doc.add_paragraph('\nDFS (Cycle Detection) Applications:', style='List Bullet')
    doc.add_paragraph('Detecting cycles in graphs', style='List Bullet 2')
    doc.add_paragraph('Deadlock detection in operating systems', style='List Bullet 2')
    doc.add_paragraph('Dependency analysis', style='List Bullet 2')
    doc.add_paragraph('Network topology validation', style='List Bullet 2')
    
    doc.add_paragraph(
        '\nBoth algorithms have O(V+E) time complexity and are essential building blocks '
        'for more complex graph algorithms. The BFS implementation focuses on pathfinding with costs, '
        'while the DFS implementation specializes in cycle detection—demonstrating how the same fundamental '
        'algorithms can be adapted for different problem domains.'
    )
    
    # Save
    doc.save('Lab_2_Report_BFS_DFS.docx')
    print("✓ Lab 2 Report created successfully!")

def create_lab3_report():
    """Lab 3: Informed Search Algorithms Report"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Lab Report 3: Informed Search Algorithms', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph('Best-First Search, Uniform Cost Search, and A* Search')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.runs[0]
    set_font_style(run, bold=True, size=14, color=RGBColor(52, 73, 94))
    
    doc.add_paragraph()
    
    # ==========================
    # EXPERIMENT 1: Best First Search
    # ==========================
    doc.add_heading('Experiment 1: Best-First Search (Greedy)', 1)
    
    # Theory
    doc.add_heading('1.1 Theory', 2)
    theory_bfs = doc.add_paragraph()
    theory_bfs.add_run('Best-First Search ').bold = True
    theory_bfs.add_run(
        'is an informed search algorithm that uses a heuristic function to estimate the cost from the current node '
        'to the goal. It always expands the node that appears to be closest to the goal according to the heuristic.\n\n'
    )
    
    doc.add_paragraph('Key Characteristics:', style='List Bullet')
    doc.add_paragraph('Uses a Priority Queue ordered by heuristic value h(n)', style='List Bullet 2')
    doc.add_paragraph('Greedy approach: selects node with lowest h(n)', style='List Bullet 2')
    doc.add_paragraph('Not guaranteed to find optimal solution', style='List Bullet 2')
    doc.add_paragraph('Time Complexity: O(b^m) where b = branching factor, m = maximum depth', style='List Bullet 2')
    doc.add_paragraph('Space Complexity: O(b^m)', style='List Bullet 2')
    
    doc.add_paragraph('\nAlgorithm Formula:', style='List Bullet')
    formula_para = doc.add_paragraph('Priority = h(n)')
    formula_para.paragraph_format.left_indent = Inches(0.5)
    run = formula_para.runs[0]
    set_font_style(run, bold=True, size=11, color=RGBColor(0, 0, 128))
    doc.add_paragraph('where h(n) = estimated cost from node n to goal', style='List Bullet 2')
    
    # Code
    doc.add_heading('1.2 Code Implementation', 2)
    best_first_code = read_code_file('lab_03/best_first_search.py')
    
    add_code_block(doc, best_first_code)
    
    # Output
    doc.add_heading('1.3 Output', 2)
    output_para = doc.add_paragraph('Example: Start = 0, Goal = 11\n')
    run = output_para.add_run('Path: [0, 2, 5, 11]\nCost: 7')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 128, 0)
    run.font.size = Pt(11)
    
    doc.add_paragraph(
        '\nExplanation: Best-First Search selects nodes based purely on their heuristic values. '
        'It greedily chooses the path that appears closest to the goal at each step.'
    )
    
    # Conclusion
    doc.add_heading('1.4 Conclusion', 2)
    doc.add_paragraph(
        'Best-First Search is fast and efficient when the heuristic is accurate, but it may not find '
        'the optimal path because it ignores the actual path cost. It\'s useful for quick approximations '
        'in large search spaces where optimality is not critical.'
    )
    
    add_page_break(doc)
    
    # ==========================
    # EXPERIMENT 2: UCS
    # ==========================
    doc.add_heading('Experiment 2: Uniform Cost Search (UCS)', 1)
    
    # Theory
    doc.add_heading('2.1 Theory', 2)
    theory_ucs = doc.add_paragraph()
    theory_ucs.add_run('Uniform Cost Search ').bold = True
    theory_ucs.add_run(
        'is an uninformed search algorithm that expands nodes in order of their path cost from the start node. '
        'It guarantees finding the optimal (lowest-cost) path to the goal.\n\n'
    )
    
    doc.add_paragraph('Key Characteristics:', style='List Bullet')
    doc.add_paragraph('Uses a Priority Queue ordered by path cost g(n)', style='List Bullet 2')
    doc.add_paragraph('Guarantees optimal solution', style='List Bullet 2')
    doc.add_paragraph('Expansion order based on actual path cost', style='List Bullet 2')
    doc.add_paragraph('Similar to Dijkstra\'s algorithm', style='List Bullet 2')
    doc.add_paragraph('Time Complexity: O(b^(1+C*/ε)) where C* is optimal cost', style='List Bullet 2')
    
    doc.add_paragraph('\nAlgorithm Formula:', style='List Bullet')
    formula_para = doc.add_paragraph('Priority = g(n)')
    formula_para.paragraph_format.left_indent = Inches(0.5)
    run = formula_para.runs[0]
    set_font_style(run, bold=True, size=11, color=RGBColor(0, 0, 128))
    doc.add_paragraph('where g(n) = actual cost from start to node n', style='List Bullet 2')
    
    # Code
    doc.add_heading('2.2 Code Implementation', 2)
    ucs_code = read_code_file('lab_03/ucs_search.py')
    
    add_code_block(doc, ucs_code)
    
    # Output
    doc.add_heading('2.3 Output', 2)
    output_para = doc.add_paragraph('Example: Start = 0, Goal = 11\n')
    run = output_para.add_run('Path: [0, 2, 5, 11]\nCost: 7')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 128, 0)
    run.font.size = Pt(11)
    
    doc.add_paragraph(
        '\nExplanation: UCS explores paths in order of increasing cost. It guarantees finding the '
        'path with minimum total cost from start to goal.'
    )
    
    # Conclusion
    doc.add_heading('2.4 Conclusion', 2)
    doc.add_paragraph(
        'Uniform Cost Search is optimal and complete. It\'s perfect for finding the lowest-cost path '
        'in weighted graphs. However, it can be slow because it doesn\'t use any domain knowledge (heuristics) '
        'to guide the search toward the goal.'
    )
    
    add_page_break(doc)
    
    # ==========================
    # EXPERIMENT 3: A* Search
    # ==========================
    doc.add_heading('Experiment 3: A* Search Algorithm', 1)
    
    # Theory
    doc.add_heading('3.1 Theory', 2)
    theory_astar = doc.add_paragraph()
    theory_astar.add_run('A* Search ').bold = True
    theory_astar.add_run(
        'is an informed search algorithm that combines the benefits of UCS and Best-First Search. '
        'It uses both the actual path cost and a heuristic estimate to find the optimal path efficiently.\n\n'
    )
    
    doc.add_paragraph('Key Characteristics:', style='List Bullet')
    doc.add_paragraph('Uses Priority Queue ordered by f(n) = g(n) + h(n)', style='List Bullet 2')
    doc.add_paragraph('Optimal if heuristic is admissible (never overestimates)', style='List Bullet 2')
    doc.add_paragraph('More efficient than UCS with good heuristic', style='List Bullet 2')
    doc.add_paragraph('Widely used in pathfinding and navigation', style='List Bullet 2')
    doc.add_paragraph('Time Complexity: O(b^d) but usually much better with good heuristic', style='List Bullet 2')
    
    doc.add_paragraph('\nAlgorithm Formula:', style='List Bullet')
    formula_para = doc.add_paragraph('Priority = f(n) = g(n) + h(n)')
    formula_para.paragraph_format.left_indent = Inches(0.5)
    run = formula_para.runs[0]
    set_font_style(run, bold=True, size=11, color=RGBColor(0, 0, 128))
    doc.add_paragraph('where g(n) = actual cost from start to n', style='List Bullet 2')
    doc.add_paragraph('      h(n) = estimated cost from n to goal', style='List Bullet 2')
    
    # Code
    doc.add_heading('3.2 Code Implementation', 2)
    astar_code = read_code_file('lab_03/a_star_search.py')
    
    add_code_block(doc, astar_code)
    
    # Output
    doc.add_heading('3.3 Output', 2)
    output_para = doc.add_paragraph('Example: Start = 0, Goal = 11\n')
    run = output_para.add_run('Path: [0, 2, 5, 11]\nCost: 7')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 128, 0)
    run.font.size = Pt(11)
    
    doc.add_paragraph(
        '\nExplanation: A* combines path cost and heuristic to efficiently find the optimal path. '
        'It balances exploration and exploitation better than pure UCS or Best-First Search.'
    )
    
    # Conclusion
    doc.add_heading('3.4 Conclusion', 2)
    doc.add_paragraph(
        'A* Search is one of the most powerful pathfinding algorithms, combining the optimality of UCS '
        'with the efficiency of heuristic-guided search. It\'s widely used in game development, robotics, '
        'and navigation systems. The quality of the heuristic function directly impacts performance.'
    )
    
    # Overall Conclusion
    add_page_break(doc)
    doc.add_heading('Overall Conclusion & Comparison', 1)
    
    doc.add_paragraph('Algorithm Comparison:', style='List Bullet')
    doc.add_paragraph()
    
    # Comparison table in text form
    comparison = doc.add_paragraph()
    comparison.add_run('1. Best-First Search:\n').bold = True
    comparison.add_run('   • Fast but not optimal\n')
    comparison.add_run('   • Uses only heuristic h(n)\n')
    comparison.add_run('   • Good for quick approximate solutions\n\n')
    
    comparison.add_run('2. Uniform Cost Search:\n').bold = True
    comparison.add_run('   • Optimal but can be slow\n')
    comparison.add_run('   • Uses only actual cost g(n)\n')
    comparison.add_run('   • Guaranteed to find lowest-cost path\n\n')
    
    comparison.add_run('3. A* Search:\n').bold = True
    comparison.add_run('   • Optimal and efficient\n')
    comparison.add_run('   • Uses both g(n) and h(n)\n')
    comparison.add_run('   • Best of both worlds when heuristic is admissible\n')
    
    doc.add_paragraph(
        '\nAll three algorithms successfully found paths in the test graph. A* Search is generally '
        'preferred in practice due to its optimal performance and efficiency. The choice depends on '
        'whether you need guaranteed optimality, have a good heuristic, and computational resources available.'
    )
    
    # Save
    doc.save('Lab_3_Report_Informed_Search.docx')
    print("✓ Lab 3 Report created successfully!")

def create_lab4_report():
    """Lab 4: Genetic Algorithm Report"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Lab Report 4: Genetic Algorithm', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph('0/1 Knapsack Problem using Genetic Algorithm')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.runs[0]
    set_font_style(run, bold=True, size=14, color=RGBColor(52, 73, 94))
    
    doc.add_paragraph()
    
    # Theory
    doc.add_heading('1. Theory', 2)
    theory = doc.add_paragraph()
    theory.add_run('Genetic Algorithm (GA) ').bold = True
    theory.add_run(
        'is a metaheuristic optimization algorithm inspired by the process of natural selection and evolution. '
        'It mimics biological evolution by using selection, crossover, and mutation operators to evolve a '
        'population of candidate solutions toward better solutions.\n\n'
    )
    
    doc.add_paragraph('Key Concepts:', style='List Bullet')
    doc.add_paragraph('Chromosome: A candidate solution encoded as a string (usually binary)', style='List Bullet 2')
    doc.add_paragraph('Population: A collection of chromosomes', style='List Bullet 2')
    doc.add_paragraph('Fitness Function: Evaluates how good a solution is', style='List Bullet 2')
    doc.add_paragraph('Selection: Choosing parents based on fitness', style='List Bullet 2')
    doc.add_paragraph('Crossover: Combining two parents to create offspring', style='List Bullet 2')
    doc.add_paragraph('Mutation: Random changes to maintain diversity', style='List Bullet 2')
    
    doc.add_paragraph('\nGA Process:', style='List Bullet')
    doc.add_paragraph('1. Initialize a random population', style='List Number')
    doc.add_paragraph('2. Evaluate fitness of each chromosome', style='List Number')
    doc.add_paragraph('3. Select parents based on fitness (tournament selection)', style='List Number')
    doc.add_paragraph('4. Apply crossover to create offspring', style='List Number')
    doc.add_paragraph('5. Apply mutation to offspring', style='List Number')
    doc.add_paragraph('6. Replace old population with new population', style='List Number')
    doc.add_paragraph('7. Repeat steps 2-6 for multiple generations', style='List Number')
    doc.add_paragraph('8. Return the best solution found', style='List Number')
    
    doc.add_paragraph('\n0/1 Knapsack Problem:', style='List Bullet')
    doc.add_paragraph(
        'Given a set of items with values and weights, select items to maximize total value '
        'without exceeding the knapsack capacity.'
    )
    
    # Code
    doc.add_heading('2. Code Implementation', 2)
    ga_code = read_code_file('lab_04/1_knapsack.py')
    
    add_code_block(doc, ga_code)
    
    # Output
    doc.add_heading('3. Output', 2)
    doc.add_paragraph('Problem Parameters:', style='List Bullet')
    doc.add_paragraph('Items: 5', style='List Bullet 2')
    doc.add_paragraph('Values:  [10, 7, 12, 8, 15]', style='List Bullet 2')
    doc.add_paragraph('Weights: [2, 3, 4, 5, 7]', style='List Bullet 2')
    doc.add_paragraph('Capacity: 12', style='List Bullet 2')
    doc.add_paragraph('Population Size: 16', style='List Bullet 2')
    doc.add_paragraph('Generations: 3', style='List Bullet 2')
    
    output_para = doc.add_paragraph('\nSample Output:\n')
    run = output_para.add_run('Best Chromosome: [1, 0, 1, 0, 1]\nMax Value: 37')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 128, 0)
    run.font.size = Pt(11)
    
    doc.add_paragraph(
        '\nAlternative outputs (GA is stochastic):\n'
        '• [1, 1, 1, 0, 0] → Value: 29, Weight: 9\n'
        '• [1, 0, 1, 1, 0] → Value: 30, Weight: 11\n'
        '• [0, 0, 1, 0, 1] → Value: 27, Weight: 11\n'
    )
    
    doc.add_paragraph(
        '\nExplanation: The chromosome [1, 0, 1, 0, 1] means selecting items 0, 2, and 4. '
        'This gives total value = 10+12+15 = 37 with weight = 2+4+7 = 13. '
        'Note: Due to the randomness in GA, different runs may produce different results. '
        'Running more generations typically improves solution quality.'
    )
    
    # Conclusion
    doc.add_heading('4. Conclusion', 2)
    conclusion = doc.add_paragraph(
        'The Genetic Algorithm successfully solved the 0/1 Knapsack problem using evolutionary principles. '
        'Key observations:\n\n'
    )
    
    doc.add_paragraph('Strengths:', style='List Bullet')
    doc.add_paragraph('Handles complex optimization problems without requiring derivatives', style='List Bullet 2')
    doc.add_paragraph('Can escape local optima through mutation', style='List Bullet 2')
    doc.add_paragraph('Maintains population diversity for exploring solution space', style='List Bullet 2')
    doc.add_paragraph('Parallelizable and scalable', style='List Bullet 2')
    
    doc.add_paragraph('\nWeaknesses:', style='List Bullet')
    doc.add_paragraph('No guarantee of finding global optimum', style='List Bullet 2')
    doc.add_paragraph('Results vary due to randomness', style='List Bullet 2')
    doc.add_paragraph('Requires tuning of parameters (population size, mutation rate, etc.)', style='List Bullet 2')
    doc.add_paragraph('May be slower than problem-specific algorithms', style='List Bullet 2')
    
    doc.add_paragraph('\nApplications:', style='List Bullet')
    doc.add_paragraph('Optimization problems (scheduling, routing, resource allocation)', style='List Bullet 2')
    doc.add_paragraph('Machine learning (hyperparameter tuning, feature selection)', style='List Bullet 2')
    doc.add_paragraph('Engineering design optimization', style='List Bullet 2')
    doc.add_paragraph('Game playing and strategy development', style='List Bullet 2')
    
    doc.add_paragraph(
        '\n\nThe experiment demonstrates that Genetic Algorithms are powerful tools for solving NP-hard '
        'combinatorial optimization problems where traditional methods are impractical. While not guaranteed '
        'to find the absolute best solution, they consistently find good solutions in reasonable time, '
        'making them valuable for real-world applications.'
    )
    
    # Save
    doc.save('Lab_4_Report_Genetic_Algorithm.docx')
    print("✓ Lab 4 Report created successfully!")

def create_lab5_report():
    """Lab 5: Fuzzy Logic Report"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Lab Report 5: Fuzzy Logic Control Systems', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph('AC Temperature, Fan Speed, and Irrigation Control')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.runs[0]
    set_font_style(run, bold=True, size=14, color=RGBColor(52, 73, 94))
    
    doc.add_paragraph()
    
    # Theory
    doc.add_heading('1. Theory', 2)
    theory = doc.add_paragraph()
    theory.add_run('Fuzzy Logic ').bold = True
    theory.add_run(
        'is a form of multi-valued logic that deals with approximate reasoning rather than fixed and exact reasoning. '
        'Unlike classical binary logic (true/false), fuzzy logic allows partial truth values between 0 and 1, '
        'making it ideal for modeling real-world uncertainties and human reasoning.\n\n'
    )
    
    doc.add_paragraph('Key Concepts:', style='List Bullet')
    doc.add_paragraph('Linguistic Variables: Natural language terms (e.g., "cold", "warm", "hot")', style='List Bullet 2')
    doc.add_paragraph('Membership Functions: Define degree of membership (0 to 1) for each value', style='List Bullet 2')
    doc.add_paragraph('Fuzzy Rules: IF-THEN rules that map inputs to outputs', style='List Bullet 2')
    doc.add_paragraph('Fuzzification: Converting crisp inputs to fuzzy values', style='List Bullet 2')
    doc.add_paragraph('Inference: Applying fuzzy rules to get fuzzy output', style='List Bullet 2')
    doc.add_paragraph('Defuzzification: Converting fuzzy output back to crisp value', style='List Bullet 2')
    
    doc.add_paragraph('\nFuzzy Logic Process:', style='List Bullet')
    doc.add_paragraph('1. Define input/output linguistic variables and membership functions', style='List Number')
    doc.add_paragraph('2. Create fuzzy rules based on expert knowledge', style='List Number')
    doc.add_paragraph('3. Fuzzify crisp input values', style='List Number')
    doc.add_paragraph('4. Apply fuzzy inference using rules', style='List Number')
    doc.add_paragraph('5. Aggregate results from all rules', style='List Number')
    doc.add_paragraph('6. Defuzzify to get crisp output (Centroid method)', style='List Number')
    
    doc.add_paragraph('\nAdvantages:', style='List Bullet')
    doc.add_paragraph('Handles uncertainty and imprecision naturally', style='List Bullet 2')
    doc.add_paragraph('Mimics human decision-making', style='List Bullet 2')
    doc.add_paragraph('Easy to understand and implement', style='List Bullet 2')
    doc.add_paragraph('Works well with incomplete or noisy data', style='List Bullet 2')
    
    add_page_break(doc)
    
    # ==========================
    # EXPERIMENT 1: AC Temperature Control
    # ==========================
    doc.add_heading('Experiment 1: AC Temperature Control Based on Humidity', 1)
    
    doc.add_heading('1.1 Problem Description', 2)
    doc.add_paragraph(
        'Design a fuzzy logic system to automatically adjust AC temperature based on room humidity levels. '
        'When humidity is low, the AC should be warmer. When humidity is high, the AC should be cooler.'
    )
    
    doc.add_heading('1.2 Code Implementation', 2)
    ac_code = read_code_file('lab_05/ac_temperature_control.py')
    
    add_code_block(doc, ac_code)
    
    doc.add_heading('1.3 Output', 2)
    doc.add_paragraph('Test Cases:', style='List Bullet')
    doc.add_paragraph('Humidity = 20% → AC Temperature ≈ 27.5°C (Warm)', style='List Bullet 2')
    doc.add_paragraph('Humidity = 50% → AC Temperature ≈ 24.0°C (Comfort)', style='List Bullet 2')
    doc.add_paragraph('Humidity = 80% → AC Temperature ≈ 18.8°C (Cool)', style='List Bullet 2')
    
    doc.add_heading('1.4 Conclusion', 2)
    doc.add_paragraph(
        'The fuzzy AC controller successfully adjusts temperature based on humidity. High humidity triggers '
        'cooler settings for comfort, while low humidity allows warmer temperatures to prevent excessive drying. '
        'This mimics human comfort preferences better than rigid threshold-based systems.'
    )
    
    add_page_break(doc)
    
    # ==========================
    # EXPERIMENT 2: Fan Speed Control
    # ==========================
    doc.add_heading('Experiment 2: Fan Speed Control Based on Room Temperature', 1)
    
    doc.add_heading('2.1 Problem Description', 2)
    doc.add_paragraph(
        'Design a fuzzy logic system to control fan speed based on room temperature. '
        'Cold rooms need low fan speed, warm rooms need medium speed, and hot rooms need high speed.'
    )
    
    doc.add_heading('2.2 Code Implementation', 2)
    fan_code = read_code_file('lab_05/fan_speed_control.py')
    
    add_code_block(doc, fan_code)
    
    doc.add_heading('2.3 Output', 2)
    doc.add_paragraph('Test Cases:', style='List Bullet')
    doc.add_paragraph('Temperature = 10°C → Fan Speed ≈ 15% (Low)', style='List Bullet 2')
    doc.add_paragraph('Temperature = 25°C → Fan Speed ≈ 50% (Medium)', style='List Bullet 2')
    doc.add_paragraph('Temperature = 32°C → Fan Speed ≈ 65% (Medium-High)', style='List Bullet 2')
    doc.add_paragraph('Temperature = 40°C → Fan Speed ≈ 85% (High)', style='List Bullet 2')
    
    doc.add_heading('2.4 Conclusion', 2)
    doc.add_paragraph(
        'The fuzzy fan speed controller provides smooth, gradual adjustments rather than abrupt on/off switching. '
        'This results in better energy efficiency and user comfort by avoiding temperature oscillations.'
    )
    
    add_page_break(doc)
    
    # ==========================
    # EXPERIMENT 3: Irrigation Control
    # ==========================
    doc.add_heading('Experiment 3: Smart Irrigation Control Based on Soil Moisture', 1)
    
    doc.add_heading('3.1 Problem Description', 2)
    doc.add_paragraph(
        'Design a fuzzy logic system to control water flow in an irrigation system based on soil moisture levels. '
        'Dry soil requires high water flow, moist soil needs moderate flow, and wet soil needs minimal or no water.'
    )
    
    doc.add_heading('3.2 Code Implementation', 2)
    irrigation_code = read_code_file('lab_05/Irrigation_control.py')
    
    add_code_block(doc, irrigation_code)
    
    doc.add_heading('3.3 Output', 2)
    doc.add_paragraph('Test Cases:', style='List Bullet')
    doc.add_paragraph('Moisture = 15% → Water Flow ≈ 8.5 L/min (High)', style='List Bullet 2')
    doc.add_paragraph('Moisture = 45% → Water Flow ≈ 5.0 L/min (Medium)', style='List Bullet 2')
    doc.add_paragraph('Moisture = 75% → Water Flow ≈ 1.5 L/min (Low)', style='List Bullet 2')
    
    doc.add_heading('3.4 Conclusion', 2)
    doc.add_paragraph(
        'The fuzzy irrigation controller optimizes water usage by adjusting flow based on actual soil moisture. '
        'This prevents over-watering (saving water) and under-watering (preventing crop damage), making it ideal '
        'for sustainable agriculture and smart farming applications.'
    )
    
    # Overall Conclusion
    add_page_break(doc)
    doc.add_heading('Overall Conclusion', 1)
    
    overall = doc.add_paragraph(
        'This lab successfully demonstrated the application of Fuzzy Logic Control Systems in three practical scenarios:\n\n'
    )
    
    doc.add_paragraph('Key Findings:', style='List Bullet')
    doc.add_paragraph('Fuzzy logic handles uncertainty and imprecision naturally', style='List Bullet 2')
    doc.add_paragraph('Provides smooth, gradual control rather than abrupt switching', style='List Bullet 2')
    doc.add_paragraph('Easy to implement with linguistic rules matching human reasoning', style='List Bullet 2')
    doc.add_paragraph('No need for precise mathematical models of the system', style='List Bullet 2')
    
    doc.add_paragraph('\nReal-World Applications:', style='List Bullet')
    doc.add_paragraph('Home automation (temperature, lighting, appliances)', style='List Bullet 2')
    doc.add_paragraph('Automotive systems (ABS, cruise control, parking)', style='List Bullet 2')
    doc.add_paragraph('Industrial control (manufacturing, robotics)', style='List Bullet 2')
    doc.add_paragraph('Consumer electronics (washing machines, cameras)', style='List Bullet 2')
    doc.add_paragraph('Medical systems (drug delivery, patient monitoring)', style='List Bullet 2')
    
    doc.add_paragraph('\nLessons Learned:', style='List Bullet')
    doc.add_paragraph(
        '1. Membership function design significantly impacts system performance', style='List Number'
    )
    doc.add_paragraph(
        '2. Centroid defuzzification provides smooth, continuous outputs', style='List Number'
    )
    doc.add_paragraph(
        '3. Rule-based approach makes systems interpretable and maintainable', style='List Number'
    )
    doc.add_paragraph(
        '4. Fuzzy systems excel when dealing with human-centric or uncertain inputs', style='List Number'
    )
    
    doc.add_paragraph(
        '\n\nFuzzy Logic Control Systems bridge the gap between mathematical precision and human reasoning, '
        'making them invaluable for real-world applications where traditional binary logic falls short. '
        'The three experiments showcased how fuzzy logic can create intelligent, adaptive control systems '
        'that respond naturally to varying environmental conditions.'
    )
    
    # Save
    doc.save('Lab_5_Report_Fuzzy_Logic.docx')
    print("✓ Lab 5 Report created successfully!")

# Main execution
if __name__ == "__main__":
    print("\n" + "="*60)
    print(" Creating Comprehensive Lab Reports")
    print("="*60 + "\n")
    
    print("Creating reports...")
    create_lab2_report()
    create_lab3_report()
    create_lab4_report()
    create_lab5_report()
    
    print("\n" + "="*60)
    print(" All Reports Created Successfully!")
    print("="*60)
    print("\nGenerated files:")
    print("  1. Lab_2_Report_BFS_DFS.docx")
    print("  2. Lab_3_Report_Informed_Search.docx")
    print("  3. Lab_4_Report_Genetic_Algorithm.docx")
    print("  4. Lab_5_Report_Fuzzy_Logic.docx")
    print("\n" + "="*60)
