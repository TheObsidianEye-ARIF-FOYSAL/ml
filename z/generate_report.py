from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
doc = Document()

# Title Page
title = doc.add_heading('Assignment on Artificial Intelligence Project', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Course Information
course_info = [
    'Course Code: CSE 3207',
    'Course Title: Artificial Intelligence',
    'Topic: AI Project on Naive Bayes Classifier',
    '',
    'Submitted By: Arif Foysal Bin Haider',
    'Roll: 2103119',
    'Section: B',
    '',
    '',
    '',
]

for info in course_info:
    p = doc.add_paragraph(info)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if info:
        for run in p.runs:
            run.font.size = Pt(12)

# Project Title
doc.add_page_break()
main_title = doc.add_heading('Machine Learning Classification Using Naive Bayes Algorithm on Iris Dataset', level=1)
main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Introduction
doc.add_heading('Introduction', level=2)
intro_text = """Machine learning is a rapidly evolving branch of artificial intelligence that enables systems to learn patterns from data and make intelligent decisions without explicit programming. Among various machine learning tasks, classification plays a crucial role in identifying the category or label of an unseen data instance based on its feature values.

The Naive Bayes classifier is a probabilistic supervised learning algorithm derived from Bayes' theorem, incorporating the simplifying assumption that input features are conditionally independent given the class label. Although this assumption is often violated in real-world data, Naive Bayes has shown impressive performance due to its mathematical simplicity, computational efficiency, and robustness.

This project presents a complete implementation of a Gaussian Naive Bayes classifier on the well-known Iris dataset, demonstrating its effectiveness in multi-class classification. Additionally, a Decision Tree classifier is implemented for comparison to evaluate relative performance and gain deeper insight into the strengths and limitations of Naive Bayes."""
doc.add_paragraph(intro_text)

# Objectives
doc.add_heading('Objectives', level=2)
objectives = [
    'To implement a Naive Bayes classifier using Python and the scikit-learn library',
    'To perform essential data preprocessing, including data validation, normalization, and dataset splitting',
    'To evaluate the classification model using performance metrics such as accuracy, precision, recall, F1-score, and confusion matrix',
    'To compare the performance of Naive Bayes with a Decision Tree classifier',
    'To analyze experimental results and explain the observed model behavior',
    'To gain practical experience in probabilistic classification and end-to-end machine learning workflow'
]
for obj in objectives:
    doc.add_paragraph(obj, style='List Bullet')

# Dataset Description
doc.add_heading('Dataset Description', level=2)

doc.add_heading('Dataset Selection', level=3)
dataset_info = """Dataset Name: Iris Dataset
Source: Kaggle
Problem Type: Multi-class classification
Data Type: Multivariate numerical data"""
doc.add_paragraph(dataset_info)

doc.add_heading('Dataset Overview', level=3)
overview_text = """The Iris dataset is a classical benchmark dataset frequently used to evaluate classification algorithms. It was introduced by statistician and biologist Ronald A. Fisher and remains widely used due to its simplicity and well-structured nature."""
doc.add_paragraph(overview_text)

# Dataset Characteristics Table
doc.add_paragraph('Dataset Characteristics:', style='Heading 4')
table1 = doc.add_table(rows=5, cols=2)
table1.style = 'Light Grid Accent 1'
characteristics = [
    ['Characteristic', 'Value'],
    ['Number of Samples', '150'],
    ['Number of Features', '4'],
    ['Number of Classes', '3'],
    ['Missing Values', 'None']
]
for i, row_data in enumerate(characteristics):
    row = table1.rows[i]
    row.cells[0].text = row_data[0]
    row.cells[1].text = row_data[1]

doc.add_paragraph()

doc.add_heading('Feature Description', level=3)
features_text = """The dataset contains four continuous numerical features measured in centimeters:
• Sepal Length – Length of the sepal
• Sepal Width – Width of the sepal
• Petal Length – Length of the petal
• Petal Width – Width of the petal"""
doc.add_paragraph(features_text)

doc.add_paragraph()
doc.add_paragraph('Statistical Summary:', style='Heading 4')

# Statistical Summary Table
table2 = doc.add_table(rows=5, cols=9)
table2.style = 'Light Grid Accent 1'
stats_data = [
    ['Feature', 'Count', 'Mean', 'Std', 'Min', '25%', '50%', '75%', 'Max'],
    ['Sepal Length', '150', '5.843', '0.828', '4.3', '5.1', '5.8', '6.4', '7.9'],
    ['Sepal Width', '150', '3.054', '0.434', '2.0', '2.8', '3.0', '3.3', '4.4'],
    ['Petal Length', '150', '3.759', '1.764', '1.0', '1.6', '4.35', '5.1', '6.9'],
    ['Petal Width', '150', '1.199', '0.763', '0.1', '0.3', '1.3', '1.8', '2.5']
]
for i, row_data in enumerate(stats_data):
    row = table2.rows[i]
    for j, cell_data in enumerate(row_data):
        row.cells[j].text = cell_data

doc.add_paragraph()

doc.add_heading('Target Classes', level=3)
classes_text = """The target variable represents three species of Iris flowers:
• Iris-setosa – 50 samples
• Iris-versicolor – 50 samples
• Iris-virginica – 50 samples

The dataset is perfectly balanced with equal representation of all three classes."""
doc.add_paragraph(classes_text)

# Methodology
doc.add_page_break()
doc.add_heading('Methodology', level=2)

doc.add_heading('Data Preprocessing', level=3)

preprocessing_steps = [
    ('Missing Value Analysis', 'The dataset was examined for missing or null values across all attributes.\nResult: No missing values were found, eliminating the need for imputation.'),
    ('Dataset Splitting', 'The dataset was divided into training and testing sets using stratified sampling to preserve class balance:\n• Training Set: 80% (120 samples)\n• Testing Set: 20% (30 samples)\nA fixed random state of 42 was used to ensure reproducibility.'),
    ('Feature Scaling', 'To ensure equal contribution of all features, StandardScaler was applied. Feature values were transformed using the formula:\n\nz = (x - μ) / σ\n\nwhere:\n• μ = mean of the feature\n• σ = standard deviation of the feature\n\nThis ensures that all features contribute equally, particularly for distance-based algorithms.')
]

for step_title, step_desc in preprocessing_steps:
    doc.add_paragraph(f'{step_title}:', style='Heading 4')
    doc.add_paragraph(step_desc)

doc.add_heading('Model Implementation', level=3)

doc.add_paragraph('Naive Bayes Classifier:', style='Heading 4')
nb_impl = """Algorithm Used: Gaussian Naive Bayes
Theoretical Basis: Bayes' Theorem

P(y|X) = P(X|y) × P(y) / P(X)

The Naive Bayes assumption considers features to be conditionally independent. For Gaussian Naive Bayes, each feature is assumed to follow a normal distribution.

The model was implemented using GaussianNB from scikit-learn. Training involved calculating the mean and variance of each feature for every class, while predictions were generated using maximum a posteriori (MAP) probability estimation."""
doc.add_paragraph(nb_impl)

# Results
doc.add_page_break()
doc.add_heading('Results', level=2)

doc.add_heading('Naive Bayes Classifier Performance', level=3)

results_text = """Test Accuracy: 96.67%
Cross-Validation Accuracy (5-fold CV): 95.83%"""
doc.add_paragraph(results_text)

doc.add_paragraph()
doc.add_paragraph('Classification Report:', style='Heading 4')

# Classification Report Table
table3 = doc.add_table(rows=5, cols=5)
table3.style = 'Light Grid Accent 1'
report_data = [
    ['Class', 'Precision', 'Recall', 'F1-Score', 'Support'],
    ['Iris-setosa', '1.00', '1.00', '1.00', '10'],
    ['Iris-versicolor', '1.00', '0.90', '0.95', '10'],
    ['Iris-virginica', '0.91', '1.00', '0.95', '10'],
    ['Overall Accuracy', '', '', '0.97', '30']
]
for i, row_data in enumerate(report_data):
    row = table3.rows[i]
    for j, cell_data in enumerate(row_data):
        row.cells[j].text = cell_data

doc.add_paragraph()

doc.add_paragraph('Confusion Matrix:', style='Heading 4')
confusion_text = """The confusion matrix shows:
                    Predicted
                 Set  Ver  Vir
Actual    Set   [10   0   0]
          Ver   [ 0   9   1]
          Vir   [ 0   0  10]

Interpretation: Only 1 misclassification occurred (Versicolor → Virginica), while all other 29 samples were classified correctly. The model achieved perfect classification for Iris-setosa and Iris-virginica classes."""
doc.add_paragraph(confusion_text)

# Decision Tree Results
doc.add_heading('Decision Tree Classifier Performance', level=3)

dt_results = """Test Accuracy: 96.67%
Cross-Validation Accuracy (5-fold CV): 93.33%
Max Depth: 3 (to reduce overfitting)"""
doc.add_paragraph(dt_results)

doc.add_paragraph()
doc.add_paragraph('Classification Report:', style='Heading 4')

# DT Classification Report Table
table4 = doc.add_table(rows=5, cols=5)
table4.style = 'Light Grid Accent 1'
dt_report_data = [
    ['Class', 'Precision', 'Recall', 'F1-Score', 'Support'],
    ['Iris-setosa', '1.00', '1.00', '1.00', '10'],
    ['Iris-versicolor', '1.00', '0.90', '0.95', '10'],
    ['Iris-virginica', '0.91', '1.00', '0.95', '10'],
    ['Overall Accuracy', '', '', '0.97', '30']
]
for i, row_data in enumerate(dt_report_data):
    row = table4.rows[i]
    for j, cell_data in enumerate(row_data):
        row.cells[j].text = cell_data

doc.add_paragraph()

# Model Comparison
doc.add_heading('Model Comparison', level=3)

# Comparison Table
table5 = doc.add_table(rows=6, cols=4)
table5.style = 'Light Grid Accent 1'
comparison_data = [
    ['Metric', 'Naive Bayes', 'Decision Tree', 'Better Model'],
    ['Test Accuracy', '96.67%', '96.67%', 'Tie'],
    ['CV Accuracy', '95.83%', '93.33%', 'Naive Bayes'],
    ['Precision', '97%', '97%', 'Tie'],
    ['Recall', '96.67%', '96.67%', 'Tie'],
    ['F1-Score', '96.66%', '96.66%', 'Tie']
]
for i, row_data in enumerate(comparison_data):
    row = table5.rows[i]
    for j, cell_data in enumerate(row_data):
        row.cells[j].text = cell_data

doc.add_paragraph()

comparison_conclusion = """While both models achieved identical test accuracy (96.67%), Naive Bayes demonstrated superior generalization with higher cross-validation accuracy (95.83% vs 93.33%). This indicates that Naive Bayes is more robust and consistent across different data splits, making it the better choice for this dataset."""
doc.add_paragraph(comparison_conclusion)

# Analysis & Discussion
doc.add_page_break()
doc.add_heading('Analysis & Discussion', level=2)

doc.add_heading('Reasons for Naive Bayes Effectiveness', level=3)

reasons = [
    ('Low Feature Correlation', 'The Iris dataset features exhibit relatively low inter-feature correlation, which aligns well with the Naive Bayes independence assumption. While features are not perfectly independent, the correlation is weak enough that the assumption does not significantly harm performance.'),
    ('Gaussian Distribution', 'The continuous features in the Iris dataset approximately follow normal (Gaussian) distributions, making Gaussian Naive Bayes a natural fit for this data. The algorithm\'s assumption that features are normally distributed per class is reasonably satisfied.'),
    ('Clear Class Separation', 'The dataset exhibits strong class separability, particularly for Iris-setosa which is linearly separable from the other two species. Even though Iris-versicolor and Iris-virginica have some overlap, the overall separation is sufficient for probabilistic classification.'),
    ('Balanced Dataset', 'With exactly 50 samples per class, the dataset is perfectly balanced. This prevents class imbalance bias and ensures that the prior probabilities are equal, allowing the model to focus on likelihood probabilities.'),
    ('Low Dimensionality', 'With only 4 features and 150 samples, the dataset has a favorable sample-to-feature ratio (37.5:1). This reduces the risk of overfitting and provides sufficient data for reliable parameter estimation.'),
    ('Simplicity and Efficiency', 'Naive Bayes requires minimal hyperparameter tuning and trains extremely fast. Its simplicity makes it less prone to overfitting compared to more complex models, especially on small datasets.')
]

for reason_title, reason_desc in reasons:
    doc.add_paragraph(f'{reason_title}:', style='Heading 4')
    doc.add_paragraph(reason_desc)

doc.add_heading('Model Limitations', level=3)

limitations = """Despite excellent performance on the Iris dataset, Naive Bayes has limitations:
• The independence assumption may not hold for datasets with highly correlated features
• Continuous features that deviate significantly from normal distribution may reduce accuracy
• Cannot learn complex feature interactions that tree-based models can capture
• Sensitive to feature scaling when using Gaussian assumption"""
doc.add_paragraph(limitations)

doc.add_heading('Key Insights', level=3)

insights = """1. Naive Bayes achieved 96.67% test accuracy with only one misclassification
2. The misclassified sample was an Iris-versicolor predicted as Iris-virginica, which is understandable given the natural overlap between these two species
3. Cross-validation accuracy (95.83%) closely matches test accuracy, indicating good model stability and no overfitting
4. The model demonstrates excellent precision and recall across all three classes
5. Naive Bayes and Decision Tree achieved identical test accuracy, but Naive Bayes showed better cross-validation performance"""
doc.add_paragraph(insights)

# Conclusion
doc.add_page_break()
doc.add_heading('Conclusion', level=2)

conclusion_text = """This project successfully implemented and evaluated a Gaussian Naive Bayes classifier on the Iris dataset, achieving 96.67% test accuracy and 95.83% cross-validation accuracy. The implementation included comprehensive data preprocessing, model training, and rigorous evaluation using multiple metrics.

Key Achievements:
• Successfully implemented Naive Bayes classifier using scikit-learn
• Performed proper data preprocessing including feature scaling and stratified splitting
• Achieved high classification accuracy with minimal misclassifications
• Conducted thorough model evaluation using accuracy, precision, recall, F1-score, and confusion matrix
• Compared Naive Bayes with Decision Tree classifier, demonstrating Naive Bayes's superior cross-validation performance
• Gained practical insights into probabilistic classification and model evaluation

The study demonstrates that simple probabilistic models like Naive Bayes can deliver excellent performance on well-structured, low-dimensional datasets with approximately Gaussian features and moderate feature independence. The high accuracy and robust cross-validation results confirm that Naive Bayes is an excellent choice for the Iris classification task.

The workflow provided valuable hands-on experience with the complete machine learning pipeline, from data preprocessing and model training to comprehensive evaluation and performance comparison. This project successfully achieved all stated objectives and delivered a reliable classification solution for the Iris dataset."""
doc.add_paragraph(conclusion_text)

# References
doc.add_page_break()
doc.add_heading('References', level=2)

references = [
    'Fisher, R.A. (1936). "The use of multiple measurements in taxonomic problems". Annals of Eugenics, 7(2), 179-188.',
    'Iris Dataset. Kaggle. Retrieved from https://www.kaggle.com/datasets/uciml/iris',
    'Pedregosa, F., et al. (2011). "Scikit-learn: Machine Learning in Python". Journal of Machine Learning Research, 12, 2825-2830.',
    'Duda, R.O., Hart, P.E., & Stork, D.G. (2001). Pattern Classification (2nd ed.). Wiley-Interscience.',
    'Murphy, K.P. (2012). Machine Learning: A Probabilistic Perspective. MIT Press.',
    'Bishop, C.M. (2006). Pattern Recognition and Machine Learning. Springer.',
    'James, G., Witten, D., Hastie, T., & Tibshirani, R. (2013). An Introduction to Statistical Learning. Springer.',
    'UCI Machine Learning Repository. (1988). Iris Data Set. Retrieved from https://archive.ics.uci.edu/ml/datasets/iris'
]

for i, ref in enumerate(references, 1):
    doc.add_paragraph(f'{i}. {ref}')

# Appendix
doc.add_page_break()
doc.add_heading('Appendix', level=2)

doc.add_heading('Code Implementation', level=3)
code_note = """The complete implementation code is available in the accompanying Jupyter Notebook file (flower.ipynb), which includes:
• Data loading and preprocessing
• Exploratory Data Analysis (EDA) with visualizations
• Feature scaling and train-test split
• Naive Bayes classifier implementation
• Decision Tree classifier for comparison
• Comprehensive model evaluation
• Confusion matrices and classification reports
• Cross-validation analysis"""
doc.add_paragraph(code_note)

doc.add_heading('Technologies Used', level=3)
tech_list = [
    'Python 3.x',
    'NumPy – Numerical computing',
    'Pandas – Data manipulation',
    'Matplotlib & Seaborn – Data visualization',
    'Scikit-learn – Machine learning algorithms',
    'Jupyter Notebook – Interactive development environment'
]
for tech in tech_list:
    doc.add_paragraph(tech, style='List Bullet')

# Save the document
output_path = r'd:\Matchine_Learning\CampusX\tutorial\z\2103119_AI_Assignment_Naive_Bayes.docx'
doc.save(output_path)
print(f"Report generated successfully: {output_path}")
